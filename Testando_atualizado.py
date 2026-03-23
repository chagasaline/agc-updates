
# -*- coding: utf-8 -*-
"""
AGC Professional Suite — Single File Build
Preserva a lógica da aba "Gerar Cargas" e adiciona melhorias não-invasivas:
- Pré-visualização (Dry-Run) antes de salvar
- Processamento em Lote (pasta)
- Auditoria (manifesto com hashes)
- Normalização opcional de termos (Dicionário Técnico)
- Redação/Compliance (PDF: remoção de CPFs, CNPJs, e-mails, telefones)
- Monitoramento de Pasta (Watcher simples)
- Relatório Executivo (PDF) pós-execução
- CLI (headless)

Requisitos (recomendados):
pip install customtkinter python-docx openpyxl matplotlib reportlab PyMuPDF pdfplumber
mas isso pode ser solicitado no txt de requisitos e deve ser visto com o administrador do sistema, para evitar dependências desnecessárias.
"""
import os, sys, json, hashlib, uuid, datetime, re, base64, time, threading, argparse, queue
from pathlib import Path

import tkinter as tk
from tkinter import filedialog, messagebox, Toplevel
import customtkinter as ctk

from docx import Document
import difflib
import openpyxl

# Matplotlib (offscreen no gráfico)
import matplotlib
matplotlib.use("Agg")
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure

# ReportLab (para relatórios PDF)
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    REPORTLAB_OK = True
except Exception:
    REPORTLAB_OK = False

# PyMuPDF para redação PDF
try:
    import fitz
    PYMUPDF_OK = True
except Exception:
    PYMUPDF_OK = False

APP_NAME = "AGC Professional Suite"
APP_VERSION = "2.1.0"  # nova build (melhorias não-invasivas)

# Pastas/arquivos persistentes
DATA_DIR         = Path.cwd() / "agc_data"
DATA_DIR.mkdir(parents=True, exist_ok=True)
USERS_PATH       = DATA_DIR / "users.json"
HISTORICO_PATH   = DATA_DIR / "historico_agc.json"
KPI_PATH         = DATA_DIR / "kpi_data.json"
CONFIG_PATH      = DATA_DIR / "config_agc.json"
LICENSE_PATH     = DATA_DIR / "license_agc.json"
AUDIT_DIR        = DATA_DIR / "auditoria"
AUDIT_DIR.mkdir(exist_ok=True)

VALIDATION_LOGS  = DATA_DIR / "logs_validacao"
VALIDATION_LOGS.mkdir(exist_ok=True)

# Validações originais do template
REQUIRED_HEADERS = [
    "Name",
    "Short description",
    "Short Description INC",
    "Configuration item",
    "Category",
    "Subcategory",
    "Impact",
    "Urgency",
    "Assignment group",
    "Description",
    "State",
    "Fill assigned to with the current user?",
    "Active",
    "Global",
    "Groups",
    "Article",
]
ALLOWED_STATES = {"Open", "Resolved", "Closed"}
ALLOWED_PRIORITIES = {"1 - Critical", "2 - High", "3 - Medium", "4 - Low"}
REQUIRED_CONSTANTS = {
    "Active": "TRUE",
    "Global": "FALSE",
    "Groups": "ITIL Processes Knowledge",
}

# -------- Utilidades --------
def _now_iso():
    return datetime.datetime.now().isoformat()

def _sha256(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1<<20), b""):
            h.update(chunk)
    return h.hexdigest()

def _write_json(path: Path, data: dict):
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

def _read_json(path: Path, default):
    if path.exists():
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            return default
    return default

# -------- Auditoria --------
def write_manifest(inputs, outputs, schema_version: str, app_version: str, notes: dict=None):
    """Cria manifesto da execução com hashes dos arquivos de entrada/saída."""
    run_dir = AUDIT_DIR / datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    run_dir.mkdir(parents=True, exist_ok=True)
    run = {
        "timestamp": datetime.datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"),
        "app_version": app_version,
        "schema_version": schema_version,
        "inputs": [{"path": str(p), "sha256": _sha256(p)} for p in inputs if Path(p).exists()],
        "outputs":[{"path": str(p), "sha256": _sha256(p)} for p in outputs if Path(p).exists()],
        "notes": notes or {}
    }
    _write_json(run_dir / "AGC_manifest.json", run)
    return str(run_dir)

# -------- Usuários/Login --------
def hash_password(password: str, salt: str = None) -> dict:
    if salt is None:
        salt = uuid.uuid4().hex
    h = hashlib.sha256((salt + password).encode("utf-8")).hexdigest()
    return {"salt": salt, "hash": h}

def verify_password(password: str, salt: str, h: str) -> bool:
    return hashlib.sha256((salt + password).encode("utf-8")).hexdigest() == h

def load_users():
    return _read_json(USERS_PATH, [])

def save_users(users):
    _write_json(USERS_PATH, users)

def ensure_admin_exists():
    users = load_users()
    if not users:
        creds = hash_password("admin123")
        users.append({
            "username": "admin",
            "name": "Administrador",
            "role": "admin",
            "password_hash": creds["hash"],
            "salt": creds["salt"],
            "created_at": _now_iso(),
        })
        save_users(users)

def registrar_historico(operacao: str, arquivo: str, status: str, usuario: str):
    hist = _read_json(HISTORICO_PATH, [])
    hist.append({
        "timestamp": _now_iso(),
        "operacao": operacao,
        "arquivo": arquivo,
        "status": status,
        "usuario": usuario,
    })
    _write_json(HISTORICO_PATH, hist)

def load_kpi_data():
    return _read_json(KPI_PATH, {"resets": [], "notes": ""})

def save_kpi_data(kpi):
    _write_json(KPI_PATH, kpi)

def reset_kpi():
    kpi = load_kpi_data()
    kpi.setdefault("resets", []).append(_now_iso())
    save_kpi_data(kpi)
    if HISTORICO_PATH.exists():
        dt = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        backup = HISTORICO_PATH.with_name(f"historico_agc_backup_{dt}.json")
        HISTORICO_PATH.rename(backup)

# -------- Normalizers (Dicionário Técnico) --------
def apply_normalizers(campos: dict, enabled: bool, rules: dict):
    """Aplica normalizações se habilitado. Não altera a lógica de extração, apenas refina valores."""
    if not enabled or not rules:
        return campos
    out = dict(campos)
    for k, v in list(out.items()):
        if isinstance(v, str):
            vv = v
            # unidades
            vv = re.sub(r"\bkgs?\b", "kg", vv, flags=re.IGNORECASE)
            vv = re.sub(r"\blitros?\b|\bltrs?\b|\bl\b", "L", vv, flags=re.IGNORECASE)
            # substituições do glossário/rules
            for src, dst in rules.items():
                vv = re.sub(rf"\b{re.escape(src)}\b", dst, vv, flags=re.IGNORECASE)
            out[k] = vv
    return out

# -------- Redação/Compliance --------
REDACT_PATTERNS = [
    r"\b\d{3}\.?\d{3}\.?\d{3}\-?\d{2}\b",         # CPF
    r"\b\d{2}\.?\d{3}\.?\d{3}/\d{4}\-?\d{2}\b",   # CNPJ
    r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
    r"\b(?:\+?55\s?)?(?:\(?\d{2}\)?\s?)?\d{4,5}\-?\d{4}\b"
]

def redact_pdf(input_path, output_path):
    if not PYMUPDF_OK:
        raise RuntimeError("PyMuPDF (fitz) não instalado.")
    doc = fitz.open(input_path)
    compiled = [re.compile(p) for p in REDACT_PATTERNS]
    for page in doc:
        blocks = page.get_text("blocks")
        for (_x0, _y0, _x1, _y1, txt, *_rest) in blocks:
            for rgx in compiled:
                for m in rgx.finditer(txt or ""):
                    rects = page.search_for(m.group())
                    for r in rects:
                        page.add_redact_annot(r, fill=(0,0,0))
        page.apply_redactions()
    doc.save(output_path, deflate=True)
    doc.close()

# -------- Classe de Login --------
class LoginWindow(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title(f"{APP_NAME} — Login")
        self.geometry("420x480")
        self.resizable(False, False)
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")

        self.current_user = None

        container = ctk.CTkFrame(self, corner_radius=12, fg_color="#232323")
        container.pack(fill="both", expand=True, padx=20, pady=20)

        title = ctk.CTkLabel(container, text="Autenticação", font=ctk.CTkFont(size=22, weight="bold"))
        title.pack(pady=(20, 10))

        self.username_var = tk.StringVar()
        self.password_var = tk.StringVar()

        user_entry = ctk.CTkEntry(container, placeholder_text="Usuário", textvariable=self.username_var, height=40)
        pass_entry = ctk.CTkEntry(container, placeholder_text="Senha", textvariable=self.password_var, height=40, show="*")
        user_entry.pack(fill="x", padx=24, pady=(8, 6))
        pass_entry.pack(fill="x", padx=24, pady=(0, 8))

        self.remember_var = tk.BooleanVar(value=True)
        remember = ctk.CTkCheckBox(container, text="Lembrar login neste dispositivo", variable=self.remember_var)
        remember.pack(anchor="w", padx=24)

        btn_login = ctk.CTkButton(container, text="Entrar", height=40, command=self.do_login)
        btn_login.pack(fill="x", padx=24, pady=(16, 8))

        ctk.CTkLabel(container, text="— ou —").pack(pady=4)

        btn_create = ctk.CTkButton(container, text="Criar novo usuário", height=36, fg_color="#0e7a0d", command=self.open_create_user)
        btn_create.pack(fill="x", padx=24, pady=(8, 8))

        self.status = ctk.CTkLabel(container, text="", text_color="#b3b3b3")
        self.status.pack(pady=(6, 16))

        ctk.CTkLabel(container, text=f"{APP_NAME} v{APP_VERSION}", text_color="#7a7a7a").pack(side="bottom", pady=8)

    def do_login(self):
        u = self.username_var.get().strip()
        p = self.password_var.get().strip()
        if not u or not p:
            self.status.configure(text="Informe usuário e senha.")
            return
        users = load_users()
        for user in users:
            if user.get("username") == u:
                if verify_password(p, user.get("salt"), user.get("password_hash")):
                    self.current_user = {"username": u, "name": user.get("name") or u, "role": user.get("role", "user")}
                    self.destroy()
                    return
                else:
                    self.status.configure(text="Senha incorreta.")
                    return
        self.status.configure(text="Usuário não encontrado.")

    def open_create_user(self):
        CreateUserDialog(self)

class CreateUserDialog(Toplevel):
    def __init__(self, parent: LoginWindow):
        super().__init__(parent)
        self.title("Criar Usuário")
        self.geometry("420x520")
        self.resizable(False, False)
        self.transient(parent)
        self.grab_set()

        frame = ctk.CTkFrame(self, corner_radius=12)
        frame.pack(fill="both", expand=True, padx=20, pady=20)

        ctk.CTkLabel(frame, text="Novo Usuário", font=ctk.CTkFont(size=20, weight="bold")).pack(pady=(16, 10))

        self.name_var = tk.StringVar()
        self.username_var = tk.StringVar()
        self.password_var = tk.StringVar()
        self.password2_var = tk.StringVar()
        self.role_var = tk.StringVar(value="user")

        ctk.CTkEntry(frame, placeholder_text="Nome completo", textvariable=self.name_var, height=38).pack(fill="x", padx=24, pady=6)
        ctk.CTkEntry(frame, placeholder_text="Usuário (login)", textvariable=self.username_var, height=38).pack(fill="x", padx=24, pady=6)
        ctk.CTkEntry(frame, placeholder_text="Senha", textvariable=self.password_var, height=38, show="*").pack(fill="x", padx=24, pady=6)
        ctk.CTkEntry(frame, placeholder_text="Confirmar senha", textvariable=self.password2_var, height=38, show="*").pack(fill="x", padx=24, pady=6)

        role_row = ctk.CTkFrame(frame, fg_color="transparent")
        role_row.pack(fill="x", padx=24, pady=8)
        ctk.CTkLabel(role_row, text="Perfil:", width=60).pack(side="left")
        ctk.CTkOptionMenu(role_row, values=["user", "admin"], variable=self.role_var, width=120).pack(side="left", padx=(8, 0))

        self.status = ctk.CTkLabel(frame, text="", text_color="#b3b3b3")
        self.status.pack(pady=(6, 8))

        actions = ctk.CTkFrame(frame, fg_color="transparent")
        actions.pack(fill="x", padx=24, pady=(8, 8))
        ctk.CTkButton(actions, text="Salvar", command=self.save_user).pack(side="left", expand=True, fill="x", padx=(0, 6))
        ctk.CTkButton(actions, text="Cancelar", fg_color="#8a1a1a", command=self.destroy).pack(side="left", expand=True, fill="x", padx=(6, 0))

    def save_user(self):
        name = self.name_var.get().strip()
        username = self.username_var.get().strip()
        password = self.password_var.get().strip()
        password2 = self.password2_var.get().strip()
        role = self.role_var.get().strip()
        if not name or not username or not password:
            self.status.configure(text="Preencha todos os campos.")
            return
        if password != password2:
            self.status.configure(text="As senhas não coincidem.")
            return
        users = load_users()
        if any(u.get("username")==username for u in users):
            self.status.configure(text="Usuário já existe.")
            return
        h = hash_password(password)
        users.append({
            "username": username,
            "name": name,
            "role": role,
            "password_hash": h["hash"],
            "salt": h["salt"],
            "created_at": _now_iso(),
        })
        save_users(users)
        messagebox.showinfo("Sucesso", "Usuário criado com sucesso!")
        self.destroy()

# -------- Aplicação Principal --------
class AGCApp:
    def __init__(self, user: dict, cli_args=None):
        self.user = user or {"username": "unknown", "name": "Usuário", "role": "user"}
        self.cli_args = cli_args or {}
        self.tema_atual = "dark"

        ctk.set_appearance_mode(self.tema_atual)
        ctk.set_default_color_theme("blue")

        self.root = ctk.CTk()
        self.root.title(APP_NAME)
        self.root.geometry("1280x820")
        self.root.minsize(1100, 720)

        self.cores = {
            "primary": "#0D6EFD",
            "secondary": "#0B5ED7",
            "accent": "#00AAFF",
            "success": "#00CC66",
            "warning": "#FF9900",
            "error": "#FF3333",
            "bg_dark": "#1A1A1A",
            "bg_card": "#2D2D2D",
            "text_primary": "#FFFFFF",
            "text_secondary": "#B3B3B3",
            "border": "#72BADB"
        }

        # Estado das melhorias (opcionais) — por padrão desligadas para não alterar comportamento
        self.normalizer_enabled = tk.BooleanVar(value=False)
        self.normalizer_rules = {}  # termo: substituição
        self.schema_version = "2.0.0"

        self._build_layout()
        self._create_pages()
        self.root.after(500, self._post_start)
        self.root.mainloop()

    # ---------- Activation Guard ----------
    def _is_activated(self) -> bool:
        if LICENSE_PATH.exists():
            try:
                data = json.loads(LICENSE_PATH.read_text(encoding="utf-8"))
                if data.get("data_ativacao") and isinstance(data.get("key",""), str) and len(data.get("key","")) >= 8:
                    return True
            except Exception:
                pass
        return True  # nesta build, não bloqueamos a UI (evita dependência externa)

    def _need_activation(self):
        messagebox.showwarning("Ativação", "Ative o produto para continuar.")
        if self.user.get("role") == "admin" and "ativacao" in self.pages:
            self.show_page("ativacao")

    def _guard(self, fn, allow_when_inactive: bool=False):
        def _wrapped(*args, **kwargs):
            if self._is_activated() or allow_when_inactive:
                return fn(*args, **kwargs)
            self._need_activation()
            return
        return _wrapped

    # ---------- Layout ----------
    def _build_layout(self):
        self.root.grid_columnconfigure(1, weight=1)
        self.root.grid_rowconfigure(0, weight=1)

        # Sidebar
        self.sidebar = ctk.CTkFrame(self.root, width=280, corner_radius=0)
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        self.sidebar.grid_rowconfigure(20, weight=1)

        # Logo e título
        title = ctk.CTkLabel(self.sidebar, text="AGC", font=ctk.CTkFont(size=34, weight="bold"), text_color=self.cores["accent"])
        subtitle = ctk.CTkLabel(self.sidebar, text="Professional Suite", font=ctk.CTkFont(size=14), text_color=self.cores["text_secondary"])
        title.pack(anchor="w", padx=20, pady=(24, 0))
        subtitle.pack(anchor="w", padx=20, pady=(0, 16))
        ctk.CTkFrame(self.sidebar, height=2, fg_color=self.cores["border"]).pack(fill="x", padx=16, pady=(0, 14))

        # Navegação
        self.nav = {}
        nav_items = [
            ("gerar", "📊  Gerar Cargas"),
            ("preview", "👁️  Pré-visualizar"),
            ("lote", "🗂️  Lote"),
            ("kpi", "📈  KPI"),
            ("validador", "✓  Validador"),
            ("compliance", "🛡️  Compliance"),
            ("automacao", "⚙️  Automação"),
            ("auditoria", "🔎  Auditoria"),
            ("config", "🧰  Configurações"),
            ("sobre", "ℹ️  Sobre"),
            ("winmerge", "🔀 WinMerge"),
            ("pdf2word", "📄 PDF → Word"),
            ("winmerge", "🔀 WinMerge"),
            ("pdf2word", "📄 PDF → Word")
        ]
        if self.user.get("role") == "admin":
            nav_items.insert(7, ("ativacao", "🔑  Ativação"))

        for key, label in nav_items:
            allow = (key == "ativacao")
            cmd = self._guard(lambda k=key: self.show_page(k), allow_when_inactive=allow)
            btn = ctk.CTkButton(self.sidebar, text=label, anchor="w", height=42,
                                fg_color="transparent", text_color=self.cores["text_secondary"],
                                hover_color=self.cores["bg_card"],
                                command=cmd)
            btn.pack(fill="x", padx=16, pady=4)
            self.nav[key] = btn

        # Info + usuário logado
        info = ctk.CTkFrame(self.sidebar, fg_color=self.cores["bg_card"])
        info.pack(fill="x", padx=16, pady=16)
        uline = f"👤 {self.user.get('name')}  •  {self.user.get('role').upper()}"
        ctk.CTkLabel(info, text=uline, font=ctk.CTkFont(size=12, weight="bold")).pack(pady=(14, 6))
        ctk.CTkLabel(info, text=f"Versão {APP_VERSION}", text_color=self.cores["text_secondary"]).pack(pady=(0, 14))

        ctk.CTkButton(self.sidebar, text="Sair", fg_color=self.cores["error"],
                      command=self._guard(self._logout, allow_when_inactive=True)).pack(fill="x", padx=16, pady=(0, 16))

        # Área principal
        self.main = ctk.CTkFrame(self.root, corner_radius=0)
        self.main.grid(row=0, column=1, sticky="nsew")
        self.main.grid_columnconfigure(0, weight=1)
        self.main.grid_rowconfigure(1, weight=1)

        # Topbar
        self.topbar = ctk.CTkFrame(self.main, height=56, corner_radius=0, fg_color=self.cores["bg_card"])
        self.topbar.grid(row=0, column=0, sticky="ew"); self.topbar.grid_columnconfigure(2, weight=1)
        ctk.CTkLabel(self.topbar, text="AGC — Automatizador de Cargas", font=ctk.CTkFont(size=18, weight="bold")).grid(row=0, column=0, padx=18, pady=12, sticky="w")
        self.lbl_user = ctk.CTkLabel(self.topbar, text=f"Usuário: {self.user.get('name')}", text_color=self.cores["text_secondary"])
        self.lbl_user.grid(row=0, column=2, padx=18, pady=12, sticky="e")

        # Content
        self.content = ctk.CTkFrame(self.main, corner_radius=0)
        self.content.grid(row=1, column=0, sticky="nsew")

        # Status bar
        self.status = ctk.CTkLabel(self.main, text="🟢 Sistema Pronto", text_color=self.cores["success"])
        self.status.grid(row=2, column=0, sticky="ew", padx=12, pady=6)

        self.pages = {}

    def _logout(self):
        self.root.destroy()
        main()  # volta ao login

    def _create_pages(self):
        self._page_gerar()
        self._page_preview()
        self._page_lote()
        self._page_kpi()
        self._page_validador()
        self._page_compliance()
        self._page_automacao()
        self._page_auditoria()
        if self.user.get("role") == "admin":
            self._page_ativacao()
        self._page_config()
        self._page_sobre()
        self._page_winmerge()
        self._page_pdf2word()
        self._page_winmerge()
        self._page_pdf2word()
        self.show_page("gerar")

    def show_page(self, key):
        for k, frame in self.pages.items():
            frame.grid_forget()
        frame = self.pages.get(key)
        if frame:
            frame.grid(row=0, column=0, sticky="nsew")
            self.content.grid_rowconfigure(0, weight=1)
            self.content.grid_columnconfigure(0, weight=1)

    def _post_start(self):
        pass

    # ---------- Página: GERAR CARGAS (lógica preservada) ----------
    def _page_gerar(self):
        page = ctk.CTkFrame(self.content, fg_color="transparent")
        self.pages["gerar"] = page

        header = ctk.CTkFrame(page, fg_color=self.cores["bg_card"], height=70)
        header.pack(fill="x", padx=16, pady=(16, 0)); header.pack_propagate(False)
        ctk.CTkLabel(header, text="📊 CARGA DE TEMPLATES | SERVICENOW", font=ctk.CTkFont(size=22, weight="bold")).pack(side="left", padx=18, pady=18)

        cfg = ctk.CTkFrame(page, fg_color=self.cores["bg_card"]); cfg.pack(fill="x", padx=16, pady=16)

        ctk.CTkLabel(cfg, text="📁 Documento de Origem (Word .docx)", font=ctk.CTkFont(size=15, weight="bold")).pack(anchor="w", padx=16, pady=(16, 8))

        row = ctk.CTkFrame(cfg, fg_color="transparent"); row.pack(fill="x", padx=16, pady=(0, 16))
        self.word_path_var = tk.StringVar()
        ctk.CTkEntry(row, textvariable=self.word_path_var, height=40).pack(side="left", fill="x", expand=True, padx=(0, 8))
        ctk.CTkButton(row, text="📂 Buscar", width=120, height=40, command=self._guard(self._pick_word)).pack(side="right")

        opts = ctk.CTkFrame(cfg, fg_color="transparent"); opts.pack(fill="x", padx=16, pady=(0,8))
        ctk.CTkCheckBox(opts, text="Aplicar Dicionário Técnico (normalizar termos)", variable=self.normalizer_enabled).pack(side="left")
        self.gloss_var = tk.StringVar(value="kg=kg\nlitro=L")
        ctk.CTkLabel(cfg, text="Glossário (termo=substituição), um por linha:", text_color=self.cores["text_secondary"]).pack(anchor="w", padx=16)
        self.gloss_box = ctk.CTkTextbox(cfg, height=70); self.gloss_box.insert("1.0", self.gloss_var.get()); self.gloss_box.pack(fill="x", padx=16, pady=(4,10))

        prog = ctk.CTkFrame(page, fg_color=self.cores["bg_card"]); prog.pack(fill="x", padx=16, pady=(0, 16))
        ctk.CTkLabel(prog, text="📈 Progresso:", font=ctk.CTkFont(size=14, weight="bold")).pack(anchor="w", padx=16, pady=(16, 8))
        self.pb = ctk.CTkProgressBar(prog, height=18, progress_color=self.cores["accent"]); self.pb.pack(fill="x", padx=16, pady=(0, 6)); self.pb.set(0)
        self.lbl_proc = ctk.CTkLabel(prog, text="Sistema pronto para processar documentos", text_color=self.cores["text_secondary"]); self.lbl_proc.pack(anchor="w", padx=16, pady=(0, 12))
        self.lbl_templates_total = ctk.CTkLabel(prog, text="Total de templates: 0", font=ctk.CTkFont(size=12, weight="bold"), text_color=self.cores["text_secondary"]); self.lbl_templates_total.pack(anchor="w", padx=16, pady=(0, 16))

        actions = ctk.CTkFrame(page, fg_color="transparent"); actions.pack(fill="x", padx=16, pady=(0, 18))
        ctk.CTkButton(actions, text="👁️ Pré-visualizar", height=44, command=self._guard(self._preview_from_word)).pack(side="left", padx=(0,8))
        ctk.CTkButton(actions, text="🚀 Processar e Gerar Planilha", height=44, fg_color=self.cores["primary"], hover_color=self.cores["secondary"], command=self._guard(self._processar_cargas_preservando_logica)).pack(side="left", padx=(0, 8))
        ctk.CTkButton(actions, text="🗑️ Limpar", height=44, fg_color=self.cores["error"], command=self._guard(self._limpar_campos_gerar)).pack(side="left")

    def _pick_word(self):
        path = filedialog.askopenfilename(title="Selecionar Documento Word", filetypes=[("Word", "*.docx")])
        if path: self.word_path_var.set(path)

    def _limpar_campos_gerar(self):
        self.word_path_var.set(""); self.pb.set(0)
        self.lbl_proc.configure(text="Sistema pronto para processar documentos")
        self.lbl_templates_total.configure(text="Total de templates: 0")

    # --------- Lógica CENTRAL preservada (extração e escrita Excel) ---------
    def _processar_cargas_preservando_logica(self):
        caminho_word = self.word_path_var.get().strip()
        if not caminho_word or not os.path.exists(caminho_word) or not caminho_word.lower().endswith(".docx"):
            messagebox.showerror("Erro", "Selecione um documento Word (.docx) válido."); return

        try:
            self.pb.set(0.05); self.lbl_proc.configure(text="Lendo o Word..."); self.root.update()
            doc = Document(caminho_word)
            texto = "\n".join(p.text for p in doc.paragraphs)
            blocos = self._extrair_blocos_templates(texto)

            if not blocos:
                messagebox.showwarning("Atenção", "Nenhum bloco 'Template Name:' encontrado.")
                self.lbl_proc.configure(text="Nenhum template encontrado."); self.pb.set(0); return

            ci_global = self._coletar_todos_ci(doc)

            self.pb.set(0.15); self.lbl_proc.configure(text="Abrindo modelo e preparando planilha..."); self.root.update()
            modelo = self._localizar_planilha_modelo()
            if not modelo or not os.path.exists(modelo):
                modelo = filedialog.askopenfilename(title="Selecionar Modelo de Templates", filetypes=[("Excel", "*.xlsx")])
                if not modelo:
                    self.lbl_proc.configure(text="Modelo não selecionado."); self.pb.set(0); return

            wb = openpyxl.load_workbook(modelo)
            ws = wb.active

            headers = [c.value for c in ws[1] if c.value]
            if ws.max_row > 1:
                ws.delete_rows(2, ws.max_row - 1)

            def escrever_linha(row_idx, campos):
                for col_idx, header in enumerate(headers, start=1):
                    ws.cell(row=row_idx, column=col_idx, value=campos.get(header, ""))

            # Normalizer rules (opcional)
            self._reload_normalizer_rules()

            total = len(blocos)
            for i, bloco in enumerate(blocos, start=1):
                campos = self._extrair_campos(bloco, ci_global)
                campos = apply_normalizers(campos, self.normalizer_enabled.get(), self.normalizer_rules)
                row = i + 1
                escrever_linha(row, campos)

                self.pb.set(0.15 + (0.70 * i / total))
                self.lbl_proc.configure(text=f"Processando template {i}/{total}...")
                self.root.update()

            self.pb.set(0.90); self.lbl_proc.configure(text="Pronto para salvar..."); self.root.update()

            destino = filedialog.asksaveasfilename(
                title="Salvar planilha gerada", defaultextension=".xlsx",
                initialfile="Modelo de templates v1.1.xlsx", filetypes=[("Excel", "*.xlsx")]
            )
            if not destino:
                self.lbl_proc.configure(text="Salvamento cancelado."); self.pb.set(0); return

            wb.save(destino)
            self.pb.set(1.0); self.lbl_proc.configure(text=f"✅ Planilha salva: {destino}")
            self.lbl_templates_total.configure(text=f"Total de templates: {total}")
            messagebox.showinfo("Sucesso", f"Planilha gerada!\nTemplates: {total}\nArquivo: {destino}")

            # ---- Auditoria (não altera lógica do Excel, apenas registra) ----
            write_manifest(inputs=[caminho_word, self._modelo_hint_for_audit()], outputs=[destino],
                           schema_version=self.schema_version, app_version=APP_VERSION,
                           notes={"mode":"single"})

            registrar_historico(operacao="Gerar Carga",
                                arquivo=os.path.basename(caminho_word),
                                status=f"Sucesso - {total} templates",
                                usuario=self.user.get("username","user"))
            # KPI não exibido aqui; fica na aba KPI
        except PermissionError as pe:
            msg = "Não foi possível salvar. O arquivo pode estar ABERTO no Excel.\nFeche-o e tente novamente."
            if "Permission denied" not in str(pe): msg = str(pe)
            messagebox.showerror("Erro de Permissão", msg)
            self.lbl_proc.configure(text="❌ Erro: acesso negado."); self.pb.set(0)
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro ao processar: {e}")
            self.lbl_proc.configure(text="❌ Erro inesperado."); self.pb.set(0)

    def _reload_normalizer_rules(self):
        txt = self.gloss_box.get("1.0","end-1c")
        rules = {}
        for line in txt.splitlines():
            if "=" in line:
                k,v = line.split("=",1)
                rules[k.strip()] = v.strip()
        self.normalizer_rules = rules

    def _modelo_hint_for_audit(self):
        nome = "Modelo de templates v1.1.xlsx"
        home = Path.home()
        candidatos = [
            Path.cwd() / nome,
            Path.cwd().parent / nome,
            home / nome,
            home / "Desktop" / nome,
            home / "Documents" / nome,
            home / "Downloads" / nome,
        ]
        for c in candidatos:
            if c.is_file():
                return str(c)
        return ""

    def _extrair_blocos_templates(self, texto: str):
        texto = texto.replace("\xa0", " ").replace(" ", " ").replace("\u200b", "")
        padrao = r"(Template\s*Name:.*?(?=Template\s*Name:|$))"
        return re.findall(padrao, texto, re.DOTALL | re.IGNORECASE)

    def _coletar_todos_ci(self, doc: Document):
        texto = "\n".join([p.text for p in doc.paragraphs])
        resultados = re.findall(r"\bCI\s*:\s*(.+)", texto, re.IGNORECASE)
        return ", ".join(valor.strip() for valor in resultados if valor.strip())

    def _extrair_campos(self, bloco: str, ci_global: str):
        campos = {
            "Name": "",
            "Short description": "",
            "Short Description INC": "",
            "Configuration item": "",
            "Category": "",
            "Subcategory": "",
            "Impact": "",
            "Urgency": "",
            "Assignment group": "",
            "Description": "",
            "State": "Open",
            "Fill assigned to with the current user?": "",
            "Active": "TRUE",
            "Global": "FALSE",
            "Groups": "ITIL Processes Knowledge",
            "Article": "",
        }
        for linha in (bloco or "").splitlines():
            if not linha.strip(): continue
            low = linha.lower()

            if low.startswith("template name:"):
                name = linha.split(":", 1)[1].strip()
                campos["Name"] = name
                up = name.upper()
                if up.startswith("L1"):
                    campos["State"] = "Resolved"
                    campos["Fill assigned to with the current user?"] = "TRUE"
                elif up.startswith("L2"):
                    campos["State"] = "Open"
                    campos["Fill assigned to with the current user?"] = "FALSE"

            elif low.startswith("short description:"):
                sd = linha.split(":", 1)[1].strip()
                campos["Short description"] = sd
                campos["Short Description INC"] = sd

            elif re.match(r"^ci\s*:", low):
                ci = linha.split(":", 1)[1].strip()
                campos["Configuration item"] = ci

            elif low.startswith("category:"):
                campos["Category"] = linha.split(":", 1)[1].strip()
            elif low.startswith("subcategory:"):
                campos["Subcategory"] = linha.split(":", 1)[1].strip()
            elif low.startswith("impact:"):
                campos["Impact"] = self._formatar_prioridade(linha.split(":", 1)[1].strip())
            elif low.startswith("urgency:"):
                campos["Urgency"] = self._formatar_prioridade(linha.split(":", 1)[1].strip())
            elif low.startswith("assignment group:"):
                campos["Assignment group"] = linha.split(":", 1)[1].strip()
            elif low.startswith("description:"):
                campos["Description"] = linha.split(":", 1)[1].strip()
            elif "kb00" in low:
                campos["Article"] = linha.strip()

        if not campos.get("Configuration item") and ci_global:
            campos["Configuration item"] = ci_global

        sd = campos.get("Short description", "").strip()
        sdi = campos.get("Short Description INC", "").strip()
        if sd and not sdi:
            campos["Short Description INC"] = sd
        elif sdi and not sd:
            campos["Short description"] = sdi
        return campos

    def _formatar_prioridade(self, valor):
        m = re.match(r"(\d)", valor or "")
        mapa = {"1": "1 - Critical", "2": "2 - High", "3": "3 - Medium", "4": "4 - Low"}
        return mapa.get(m.group(1) if m else "", "3 - Medium")

    def _localizar_planilha_modelo(self):
        nome = "Modelo de templates v1.1.xlsx"
        home = Path.home()
        candidatos = [
            Path.cwd() / nome,
            Path.cwd().parent / nome,
            home / nome,
            home / "Desktop" / nome,
            home / "Documents" / nome,
            home / "Downloads" / nome,
        ]
        for c in candidatos:
            if c.is_file():
                return str(c)
        for p in home.glob("OneDrive*/**/" + nome):
            if p.is_file():
                return str(p)
        return None

    # ---------- Página: PREVIEW (Dry-Run) ----------
    def _page_preview(self):
        page = ctk.CTkFrame(self.content, fg_color="transparent")
        self.pages["preview"] = page
        header = ctk.CTkFrame(page, fg_color=self.cores["bg_card"], height=62); header.pack(fill="x", padx=16, pady=(16, 0)); header.pack_propagate(False)
        ctk.CTkLabel(header, text="👁️ Pré-visualização (Dry-Run)", font=ctk.CTkFont(size=20, weight="bold")).pack(side="left", padx=18, pady=16)
        body = ctk.CTkFrame(page, fg_color=self.cores["bg_card"]); body.pack(fill="both", expand=True, padx=16, pady=12)
        self.preview_box = ctk.CTkTextbox(body, height=520, font=ctk.CTkFont(size=12, family="Consolas"))
        self.preview_box.pack(fill="both", expand=True, padx=12, pady=12)
        ctk.CTkButton(page, text="Atualizar Pré-visualização", command=self._guard(self._preview_from_word)).pack(pady=(0,16))

    def _preview_from_word(self):
        path = self.word_path_var.get().strip()
        if not path or not os.path.exists(path):
            messagebox.showwarning("Pré-visualização", "Selecione um documento Word válido na aba 'Gerar'."); return
        try:
            doc = Document(path)
            texto = "\n".join(p.text for p in doc.paragraphs)
            blocos = self._extrair_blocos_templates(texto)
            ci_global = self._coletar_todos_ci(doc)
            self._reload_normalizer_rules()
            lines = []
            for i, bloco in enumerate(blocos, 1):
                campos = self._extrair_campos(bloco, ci_global)
                campos_norm = apply_normalizers(campos, self.normalizer_enabled.get(), self.normalizer_rules)
                lines.append(f"--- Template #{i} ---")
                for k in REQUIRED_HEADERS:
                    if k in campos_norm:
                        lines.append(f"{k}: {campos_norm.get(k,'')}")
                lines.append("")
            txt = "\n".join(lines) if lines else "Nenhum template encontrado."
            if "preview" in self.pages:
                self.preview_box.delete("1.0", "end")
                self.preview_box.insert("1.0", txt)
            else:
                messagebox.showinfo("Pré-visualização", txt[:4000])
        except Exception as e:
            messagebox.showerror("Pré-visualização", f"Erro: {e}")

    # ---------- Página: LOTE ----------
    def _page_lote(self):
        page = ctk.CTkFrame(self.content, fg_color="transparent")
        self.pages["lote"] = page
        header = ctk.CTkFrame(page, fg_color=self.cores["bg_card"], height=62); header.pack(fill="x", padx=16, pady=(16, 0)); header.pack_propagate(False)
        ctk.CTkLabel(header, text="🗂️ Processamento em Lote", font=ctk.CTkFont(size=20, weight="bold")).pack(side="left", padx=18, pady=16)
        body = ctk.CTkFrame(page, fg_color=self.cores["bg_card"]); body.pack(fill="both", expand=True, padx=16, pady=12)

        row = ctk.CTkFrame(body, fg_color="transparent"); row.pack(fill="x", padx=16, pady=(16, 8))
        self.lote_dir_var = tk.StringVar()
        ctk.CTkEntry(row, textvariable=self.lote_dir_var, height=40).pack(side="left", fill="x", expand=True, padx=(0, 8))
        ctk.CTkButton(row, text="📁 Pasta de .docx", command=self._guard(self._pick_lote_dir)).pack(side="left")

        row2 = ctk.CTkFrame(body, fg_color="transparent"); row2.pack(fill="x", padx=16, pady=(0, 8))
        self.lote_modelo_var = tk.StringVar()
        ctk.CTkEntry(row2, textvariable=self.lote_modelo_var, height=40, placeholder_text="(opcional) Modelo .xlsx (vazio para auto-localizar)").pack(side="left", fill="x", expand=True, padx=(0, 8))
        ctk.CTkButton(row2, text="📄 Selecionar Modelo", command=self._guard(self._pick_lote_modelo)).pack(side="left")

        actions = ctk.CTkFrame(body, fg_color="transparent"); actions.pack(fill="x", padx=16, pady=(0, 8))
        ctk.CTkButton(actions, text="▶️ Iniciar Lote", fg_color=self.cores["primary"], command=self._guard(self._run_lote)).pack(side="left", padx=6)
        self.lote_log = ctk.CTkTextbox(body, height=420, font=ctk.CTkFont(size=12, family="Consolas")); self.lote_log.pack(fill="both", expand=True, padx=16, pady=12)

    def _pick_lote_dir(self):
        p = filedialog.askdirectory(title="Selecionar Pasta com .docx")
        if p: self.lote_dir_var.set(p)

    def _pick_lote_modelo(self):
        p = filedialog.askopenfilename(title="Selecionar Modelo de Templates", filetypes=[("Excel", "*.xlsx")])
        if p: self.lote_modelo_var.set(p)

    def _run_lote(self):
        folder = self.lote_dir_var.get().strip()
        if not folder or not os.path.isdir(folder):
            messagebox.showwarning("Lote", "Selecione uma pasta válida."); return
        modelo = self.lote_modelo_var.get().strip() or self._localizar_planilha_modelo()
        if not modelo or not os.path.exists(modelo):
            messagebox.showwarning("Lote", "Modelo não encontrado."); return

        files = [str(p) for p in Path(folder).glob("*.docx")]
        if not files:
            messagebox.showinfo("Lote", "Nenhum .docx encontrado."); return

        ok, fail = 0, 0
        out_paths = []
        for idx, caminho_word in enumerate(files, 1):
            try:
                # mesma lógica da geração single (ler/parse/extrair/escrever)
                doc = Document(caminho_word)
                texto = "\n".join(p.text for p in doc.paragraphs)
                blocos = self._extrair_blocos_templates(texto)
                if not blocos:
                    self.lote_log.insert("end", f"[AVISO] {os.path.basename(caminho_word)}: nenhum template.\n"); self.lote_log.see("end"); continue
                ci_global = self._coletar_todos_ci(doc)

                wb = openpyxl.load_workbook(modelo)
                ws = wb.active
                headers = [c.value for c in ws[1] if c.value]
                if ws.max_row > 1: ws.delete_rows(2, ws.max_row - 1)

                self._reload_normalizer_rules()

                for i, bloco in enumerate(blocos, start=1):
                    campos = self._extrair_campos(bloco, ci_global)
                    campos = apply_normalizers(campos, self.normalizer_enabled.get(), self.normalizer_rules)
                    row = i + 1
                    for col_idx, header in enumerate(headers, start=1):
                        ws.cell(row=row, column=col_idx, value=campos.get(header, ""))

                out = str(Path(folder) / (Path(caminho_word).stem + "_modelo.xlsx"))
                wb.save(out)
                out_paths.append(out)
                ok += 1
                self.lote_log.insert("end", f"[OK] {os.path.basename(caminho_word)} → {os.path.basename(out)}\n")
                self.lote_log.see("end")
                registrar_historico("Gerar Carga (lote)", os.path.basename(caminho_word), f"Sucesso - {len(blocos)} templates", self.user.get("username","user"))
            except Exception as e:
                fail += 1
                self.lote_log.insert("end", f"[ERRO] {os.path.basename(caminho_word)}: {e}\n")
                self.lote_log.see("end")

        # Auditoria do lote
        try:
            write_manifest(inputs=files+[modelo], outputs=out_paths, schema_version=self.schema_version, app_version=APP_VERSION, notes={"mode":"batch"})
        except Exception:
            pass
        messagebox.showinfo("Lote", f"Concluído. OK: {ok}  |  Erros: {fail}")

    # ---------- Página: KPI (resumo simples) ----------
    def _page_kpi(self):
        page = ctk.CTkFrame(self.content, fg_color="transparent")
        self.pages["kpi"] = page
        header = ctk.CTkFrame(page, fg_color=self.cores["bg_card"], height=70); header.pack(fill="x", padx=16, pady=(16,0)); header.pack_propagate(False)
        ctk.CTkLabel(header, text="📈 KPI de Templates", font=ctk.CTkFont(size=22, weight="bold")).pack(side="left", padx=18, pady=18)

        body = ctk.CTkFrame(page, fg_color=self.cores["bg_card"]); body.pack(fill="both", expand=True, padx=16, pady=12)

        self.fig = Figure(figsize=(10, 4.4), dpi=100, facecolor="#2D2D2D")
        self.ax = self.fig.add_subplot(111, facecolor="#2D2D2D")
        self.canvas = FigureCanvasTkAgg(self.fig, master=body)
        self.canvas.get_tk_widget().pack(fill="both", expand=True, padx=8, pady=8)

        ctk.CTkButton(page, text="🔄 Atualizar", command=self._guard(self._kpi_refresh)).pack(pady=(0,10))
        if REPORTLAB_OK:
            ctk.CTkButton(page, text="📑 Exportar PDF (Executivo)", command=self._guard(self._kpi_export_pdf)).pack(pady=(0,14))

        self._kpi_refresh()

    def _kpi_load_hist(self):
        return _read_json(HISTORICO_PATH, [])

    def _kpi_refresh(self):
        hist = self._kpi_load_hist()
        dados = [h for h in hist if str(h.get("operacao","")).lower().startswith("gerar carga")]
        por_data = {}
        total = 0
        for d in dados:
            mt = re.search(r"(\d+)\s*templates?", str(d.get("status","")).lower())
            count = int(mt.group(1)) if mt else 1
            dt = datetime.datetime.fromisoformat(d["timestamp"]).date()
            por_data.setdefault(dt, 0); por_data[dt] += count; total += count

        self.ax.clear()
        if por_data:
            xs = sorted(por_data.keys())
            ys = [por_data[k] for k in xs]
            self.ax.plot(xs, ys, marker="o")
            self.ax.fill_between(xs, ys, alpha=0.18)
            self.ax.set_title(f"Templates por Dia (Total: {total})", color="white")
            self.ax.set_xlabel("Data", color="white"); self.ax.set_ylabel("Qtde", color="white")
            self.ax.tick_params(colors="white")
            for spine in self.ax.spines.values():
                spine.set_color("#777777")
        else:
            self.ax.text(0.5, 0.5, "Sem dados", color="white", ha="center", va="center")
        self.fig.tight_layout(); self.canvas.draw()

    def _kpi_export_pdf(self):
        save = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF", "*.pdf")])
        if not save: return
        hist = self._kpi_load_hist()
        dados = [h for h in hist if str(h.get("operacao","")).lower().startswith("gerar carga")]
        doc = SimpleDocTemplate(str(save), pagesize=A4)
        styles = getSampleStyleSheet()
        elems = [Paragraph("Relatório KPI — AGC", styles["Title"]), Spacer(1, 12)]
        data = [["Data/Hora", "Arquivo", "Status", "Usuário", "Templates"]]
        for d in dados:
            mt = re.search(r"(\d+)\s*templates?", str(d.get("status","")).lower())
            count = int(mt.group(1)) if mt else 1
            data.append([d.get("timestamp","")[:19].replace("T"," "),
                         d.get("arquivo",""), d.get("status",""),
                         d.get("usuario",""), str(count)])
        tbl = Table(data, repeatRows=1)
        tbl.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0), colors.HexColor("#0D6EFD")),
                                 ("TEXTCOLOR",(0,0),(-1,0), colors.white),
                                 ("GRID",(0,0),(-1,-1), 0.25, colors.grey),
                                 ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold")]))
        elems.append(tbl)
        doc.build(elems)
        messagebox.showinfo("PDF", "Relatório exportado com sucesso.")

    # ---------- Página: VALIDADOR ----------
    def _page_validador(self):
        page = ctk.CTkFrame(self.content, fg_color="transparent")
        self.pages["validador"] = page
        header = ctk.CTkFrame(page, fg_color=self.cores["bg_card"], height=62); header.pack(fill="x", padx=16, pady=(16, 0)); header.pack_propagate(False)
        ctk.CTkLabel(header, text="✓ Validador de Template (Excel)", font=ctk.CTkFont(size=20, weight="bold")).pack(side="left", padx=18, pady=16)

        body = ctk.CTkFrame(page, fg_color=self.cores["bg_card"]); body.pack(fill="both", expand=True, padx=16, pady=12)
        row = ctk.CTkFrame(body, fg_color="transparent"); row.pack(fill="x", padx=16, pady=16)
        self.val_path_var = tk.StringVar()
        ctk.CTkEntry(row, textvariable=self.val_path_var, height=40).pack(side="left", fill="x", expand=True, padx=(0,8))
        ctk.CTkButton(row, text="📂 Buscar Excel", command=self._guard(self._pick_excel)).pack(side="left")
        ctk.CTkButton(row, text="🔍 Validar", fg_color=self.cores["primary"], command=self._guard(self._validar_template)).pack(side="left", padx=8)
        self.txt_val = ctk.CTkTextbox(body, height=360, font=ctk.CTkFont(size=12, family="Consolas")); self.txt_val.pack(fill="both", expand=True, padx=16, pady=(0, 12))
        ctk.CTkButton(body, text="💾 Exportar Log", command=self._guard(self._exportar_log_validacao)).pack(pady=(0,12))

    def _pick_excel(self):
        p = filedialog.askopenfilename(title="Selecionar Planilha", filetypes=[("Excel","*.xlsx")])
        if p: self.val_path_var.set(p)

    def _validar_template(self):
        self.txt_val.delete("1.0","end")
        path = self.val_path_var.get().strip()
        if not path or not os.path.exists(path):
            self.txt_val.insert("1.0","Arquivo inválido.\n"); return
        try:
            wb = openpyxl.load_workbook(path)
            ws = wb.active
            headers = [c.value for c in ws[1] if c.value]
            logs = []
            missing = [h for h in REQUIRED_HEADERS if h not in headers]
            if missing: logs.append(f"[ERRO] Cabeçalhos ausentes: {', '.join(missing)}")

            for row in ws.iter_rows(min_row=2, values_only=True):
                row_data = {headers[i]: (row[i] if i < len(row) else "") for i in range(len(headers))}
                if not row_data.get("Name"):
                    logs.append("[ERRO] 'Name' não informado.")
                st = str(row_data.get("State","")).strip() or "Open"
                if st not in ALLOWED_STATES:
                    logs.append(f"[ALERTA] State '{st}' fora do permitido {ALLOWED_STATES}.")
                for k in ("Impact","Urgency"):
                    val = str(row_data.get(k,"")).strip()
                    if val and val not in ALLOWED_PRIORITIES:
                        logs.append(f"[ALERTA] {k} '{val}' não mapeado em {ALLOWED_PRIORITIES}.")
                for k, v in REQUIRED_CONSTANTS.items():
                    if str(row_data.get(k,"")).strip() != v:
                        logs.append(f"[ERRO] Campo '{k}' deve ser '{v}'.")

            if not logs: logs.append("[OK] Nenhum problema encontrado. Template consistente.")
            self._last_validation_log = "\n".join(logs)
            self.txt_val.insert("1.0", self._last_validation_log + "\n")
        except Exception as e:
            self.txt_val.insert("1.0", f"Erro na validação: {e}\n")

    def _exportar_log_validacao(self):
        if not hasattr(self, "_last_validation_log"):
            messagebox.showinfo("Log", "Nenhuma validação realizada."); return
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        p = VALIDATION_LOGS / f"validacao_{ts}.txt"
        p.write_text(self._last_validation_log, encoding="utf-8")
        messagebox.showinfo("Log", f"Log salvo em:\n{p}")

    # ---------- Página: COMPLIANCE (Redação) ----------
    def _page_compliance(self):
        page = ctk.CTkFrame(self.content, fg_color="transparent")
        self.pages["compliance"] = page
        header = ctk.CTkFrame(page, fg_color=self.cores["bg_card"], height=62); header.pack(fill="x", padx=16, pady=(16, 0)); header.pack_propagate(False)
        ctk.CTkLabel(header, text="🛡️ Compliance (Redação de PDF)", font=ctk.CTkFont(size=20, weight="bold")).pack(side="left", padx=18, pady=16)

        body = ctk.CTkFrame(page, fg_color=self.cores["bg_card"]); body.pack(fill="both", expand=True, padx=16, pady=12)
        row = ctk.CTkFrame(body, fg_color="transparent"); row.pack(fill="x", padx=16, pady=16)
        self.pdf_in_var = tk.StringVar(); self.pdf_out_var = tk.StringVar()
        ctk.CTkEntry(row, textvariable=self.pdf_in_var, placeholder_text="PDF de entrada", height=40).pack(side="left", fill="x", expand=True, padx=(0,8))
        ctk.CTkButton(row, text="📂 PDF", command=self._guard(lambda: self._pick_file_to_var(self.pdf_in_var, [("PDF","*.pdf")]))).pack(side="left")
        row2 = ctk.CTkFrame(body, fg_color="transparent"); row2.pack(fill="x", padx=16, pady=8)
        ctk.CTkEntry(row2, textvariable=self.pdf_out_var, placeholder_text="PDF de saída (redigido)", height=40).pack(side="left", fill="x", expand=True, padx=(0,8))
        ctk.CTkButton(row2, text="💾 Salvar como", command=self._guard(lambda: self._pick_save_to_var(self.pdf_out_var, ".pdf"))).pack(side="left")
        ctk.CTkButton(body, text="🛡️ Redigir", fg_color=self.cores["primary"], command=self._guard(self._do_redact)).pack(pady=(4,12))

        self.lbl_compliance = ctk.CTkLabel(body, text="", text_color=self.cores["text_secondary"])
        self.lbl_compliance.pack(pady=(4,12))

    def _pick_file_to_var(self, var: tk.StringVar, types):
        p = filedialog.askopenfilename(filetypes=types)
        if p: var.set(p)

    def _pick_save_to_var(self, var: tk.StringVar, ext: str):
        p = filedialog.asksaveasfilename(defaultextension=ext)
        if p: var.set(p)

    def _do_redact(self):
        i, o = self.pdf_in_var.get().strip(), self.pdf_out_var.get().strip()
        if not i or not o:
            messagebox.showwarning("Compliance", "Informe os caminhos de entrada e saída."); return
        try:
            redact_pdf(i, o)
            self.lbl_compliance.configure(text=f"✅ Redação concluída: {o}")
            messagebox.showinfo("Compliance", f"PDF redigido com sucesso:\n{o}")
        except Exception as e:
            messagebox.showerror("Compliance", f"Erro: {e}")

    # ---------- Página: AUTOMAÇÃO (Watcher simples) ----------
    def _page_automacao(self):
        page = ctk.CTkFrame(self.content, fg_color="transparent")
        self.pages["automacao"] = page
        header = ctk.CTkFrame(page, fg_color=self.cores["bg_card"], height=62); header.pack(fill="x", padx=16, pady=(16, 0)); header.pack_propagate(False)
        ctk.CTkLabel(header, text="⚙️ Automação (Monitor de Pasta)", font=ctk.CTkFont(size=20, weight="bold")).pack(side="left", padx=18, pady=16)

        body = ctk.CTkFrame(page, fg_color=self.cores["bg_card"]); body.pack(fill="both", expand=True, padx=16, pady=12)
        row = ctk.CTkFrame(body, fg_color="transparent"); row.pack(fill="x", padx=16, pady=8)
        self.auto_dir_var = tk.StringVar(); self.auto_out_dir_var = tk.StringVar()
        ctk.CTkEntry(row, textvariable=self.auto_dir_var, placeholder_text="Pasta monitorada (.docx)", height=40).pack(side="left", fill="x", expand=True, padx=(0,8))
        ctk.CTkButton(row, text="📁 Selecionar", command=self._guard(lambda: self._pick_dir_to_var(self.auto_dir_var))).pack(side="left")
        row2 = ctk.CTkFrame(body, fg_color="transparent"); row2.pack(fill="x", padx=16, pady=8)
        ctk.CTkEntry(row2, textvariable=self.auto_out_dir_var, placeholder_text="Pasta de saída (.xlsx)", height=40).pack(side="left", fill="x", expand=True, padx=(0,8))
        ctk.CTkButton(row2, text="📁 Selecionar Saída", command=self._guard(lambda: self._pick_dir_to_var(self.auto_out_dir_var))).pack(side="left")
        self.auto_running = False
        ctk.CTkButton(body, text="▶️ Iniciar", fg_color=self.cores["primary"], command=self._guard(self._auto_start)).pack(side="left", padx=18, pady=8)
        ctk.CTkButton(body, text="⏹️ Parar", fg_color=self.cores["error"], command=self._guard(self._auto_stop)).pack(side="left", padx=8, pady=8)
        self.auto_log = ctk.CTkTextbox(body, height=360, font=ctk.CTkFont(size=12, family="Consolas")); self.auto_log.pack(fill="both", expand=True, padx=16, pady=12)

    def _pick_dir_to_var(self, var: tk.StringVar):
        p = filedialog.askdirectory()
        if p: var.set(p)

    def _auto_start(self):
        if self.auto_running: return
        watch_dir = self.auto_dir_var.get().strip()
        out_dir = self.auto_out_dir_var.get().strip()
        if not watch_dir or not os.path.isdir(watch_dir):
            messagebox.showwarning("Automação", "Pasta monitorada inválida."); return
        if not out_dir or not os.path.isdir(out_dir):
            messagebox.showwarning("Automação", "Pasta de saída inválida."); return
        self.auto_running = True
        self.auto_log.insert("end", "[INFO] Monitoramento iniciado.\n"); self.auto_log.see("end")
        self._auto_thread = threading.Thread(target=self._auto_loop, args=(watch_dir, out_dir), daemon=True)
        self._auto_thread.start()

    def _auto_stop(self):
        self.auto_running = False
        self.auto_log.insert("end", "[INFO] Monitoramento interrompido.\n"); self.auto_log.see("end")

    def _auto_loop(self, watch_dir, out_dir):
        seen = set()
        while self.auto_running:
            for p in Path(watch_dir).glob("*.docx"):
                sp = str(p.resolve())
                if sp in seen: continue
                seen.add(sp)
                try:
                    doc = Document(sp)
                    texto = "\n".join(px.text for px in doc.paragraphs)
                    blocos = self._extrair_blocos_templates(texto)
                    if not blocos:
                        self.auto_log.insert("end", f"[AVISO] {p.name}: nenhum template.\n"); self.auto_log.see("end"); continue
                    ci_global = self._coletar_todos_ci(doc)

                    modelo = self._localizar_planilha_modelo()
                    if not modelo: 
                        self.auto_log.insert("end", f"[ERRO] Modelo não encontrado.\n"); self.auto_log.see("end"); continue
                    wb = openpyxl.load_workbook(modelo)
                    ws = wb.active
                    headers = [c.value for c in ws[1] if c.value]
                    if ws.max_row > 1: ws.delete_rows(2, ws.max_row - 1)

                    self._reload_normalizer_rules()

                    for i, bloco in enumerate(blocos, start=1):
                        campos = self._extrair_campos(bloco, ci_global)
                        campos = apply_normalizers(campos, self.normalizer_enabled.get(), self.normalizer_rules)
                        row = i + 1
                        for col_idx, header in enumerate(headers, start=1):
                            ws.cell(row=row, column=col_idx, value=campos.get(header, ""))

                    out = str(Path(out_dir) / (p.stem + "_modelo.xlsx"))
                    wb.save(out)
                    write_manifest(inputs=[sp, self._modelo_hint_for_audit()], outputs=[out], schema_version=self.schema_version, app_version=APP_VERSION, notes={"mode":"watch"})
                    registrar_historico("Gerar Carga (watch)", p.name, f"Sucesso - {len(blocos)} templates", self.user.get("username","user"))
                    self.auto_log.insert("end", f"[OK] {p.name} → {Path(out).name}\n"); self.auto_log.see("end")
                except Exception as e:
                    self.auto_log.insert("end", f"[ERRO] {p.name}: {e}\n"); self.auto_log.see("end")
            time.sleep(2.0)

    # ---------- Página: AUDITORIA ----------
    def _page_auditoria(self):
        page = ctk.CTkFrame(self.content, fg_color="transparent")
        self.pages["auditoria"] = page
        header = ctk.CTkFrame(page, fg_color=self.cores["bg_card"], height=62); header.pack(fill="x", padx=16, pady=(16, 0)); header.pack_propagate(False)
        ctk.CTkLabel(header, text="🔎 Auditoria (Manifests)", font=ctk.CTkFont(size=20, weight="bold")).pack(side="left", padx=18, pady=16)
        body = ctk.CTkFrame(page, fg_color=self.cores["bg_card"]); body.pack(fill="both", expand=True, padx=16, pady=12)
        self.audit_list = ctk.CTkTextbox(body, height=520, font=ctk.CTkFont(size=12, family="Consolas")); self.audit_list.pack(fill="both", expand=True, padx=12, pady=12)
        ctk.CTkButton(page, text="Atualizar", command=self._guard(self._audit_refresh)).pack(pady=(0,16))
        self._audit_refresh()

    def _audit_refresh(self):
        self.audit_list.delete("1.0", "end")
        if not AUDIT_DIR.exists():
            self.audit_list.insert("1.0", "Sem auditorias.\n"); return
        runs = sorted([p for p in AUDIT_DIR.iterdir() if p.is_dir()], reverse=True)
        if not runs:
            self.audit_list.insert("1.0", "Sem auditorias.\n"); return
        lines = []
        for r in runs[:200]:
            mf = r / "AGC_manifest.json"
            if mf.exists():
                try:
                    data = json.loads(mf.read_text(encoding="utf-8"))
                    lines.append(f"[{r.name}] {data.get('timestamp','')} — {data.get('notes',{}).get('mode','single')}")
                    for it in data.get("inputs", []):
                        lines.append(f"  IN: {it.get('path')}  #{it.get('sha256','')[:12]}")
                    for ot in data.get("outputs", []):
                        lines.append(f"  OUT: {ot.get('path')} #{ot.get('sha256','')[:12]}")
                except Exception as e:
                    lines.append(f"[{r.name}] erro ao ler manifesto: {e}")
        self.audit_list.insert("1.0", "\n".join(lines))

    # ---------- Ativação/Config/Sobre ----------
    def _page_ativacao(self):
        page = ctk.CTkFrame(self.content, fg_color="transparent")
        self.pages["ativacao"] = page
        header = ctk.CTkFrame(page, fg_color=self.cores["bg_card"], height=62); header.pack(fill="x", padx=16, pady=(16, 0)); header.pack_propagate(False)
        ctk.CTkLabel(header, text="🔑 Ativação", font=ctk.CTkFont(size=20, weight="bold")).pack(side="left", padx=18, pady=16)
        body = ctk.CTkFrame(page, fg_color=self.cores["bg_card"]); body.pack(fill="both", expand=True, padx=16, pady=12)
        ctk.CTkLabel(body, text="Insira a chave para ativar o produto.").pack(pady=12)
        self.lic_var = tk.StringVar()
        ctk.CTkEntry(body, textvariable=self.lic_var, placeholder_text="SUA-CHAVE-AQUI").pack(pady=6)
        ctk.CTkButton(body, text="Ativar", command=self._guard(self._ativar_produto, allow_when_inactive=True)).pack(pady=6)

    def _ativar_produto(self):
        key = self.lic_var.get().strip()
        if not key or len(key) < 8:
            messagebox.showwarning("Ativação", "Chave inválida."); return
        data = {"key": key, "data_ativacao": _now_iso()}
        _write_json(LICENSE_PATH, data)
        messagebox.showinfo("Ativação", "Produto ativado.")
        self.show_page("gerar")

    def _page_config(self):
        page = ctk.CTkFrame(self.content, fg_color="transparent")
        self.pages["config"] = page
        header = ctk.CTkFrame(page, fg_color=self.cores["bg_card"], height=62); header.pack(fill="x", padx=16, pady=(16, 0)); header.pack_propagate(False)
        ctk.CTkLabel(header, text="🧰 Configurações", font=ctk.CTkFont(size=20, weight="bold")).pack(side="left", padx=18, pady=16)
        body = ctk.CTkFrame(page, fg_color=self.cores["bg_card"]); body.pack(fill="both", expand=True, padx=16, pady=12)

        ctk.CTkLabel(body, text="Dicionário Técnico (normalizadores): termo=valor, um por linha").pack(anchor="w", padx=16, pady=(8,4))
        self.cfg_gloss = ctk.CTkTextbox(body, height=120); self.cfg_gloss.pack(fill="x", padx=16, pady=(0,8))
        ctk.CTkButton(body, text="Salvar Glossário", command=self._guard(self._cfg_save_gloss)).pack(padx=16, pady=(0,12))

    def _cfg_save_gloss(self):
        rules = self.cfg_gloss.get("1.0","end-1c")
        CONFIG_PATH.write_text(json.dumps({"gloss": rules}, ensure_ascii=False, indent=2), encoding="utf-8")
        messagebox.showinfo("Config", "Glossário salvo.")


    # ---------- Página: WINMERGE (Comparação estilo WinMerge) ----------
    def _page_winmerge(self):
        page = ctk.CTkFrame(self.content, fg_color="transparent")
        self.pages["winmerge"] = page
        header = ctk.CTkFrame(page, fg_color=self.cores["bg_card"], height=62)
        header.pack(fill="x", padx=16, pady=(16,0)); header.pack_propagate(False)
        ctk.CTkLabel(header, text="🔀 Comparação estilo WinMerge", font=ctk.CTkFont(size=20, weight="bold")).pack(side="left", padx=18, pady=16)

        body = ctk.CTkFrame(page, fg_color=self.cores["bg_card"]); body.pack(fill="both", expand=True, padx=16, pady=12)
        self.txtA = ctk.CTkTextbox(body, width=500, height=400); self.txtA.pack(side="left", fill="both", expand=True, padx=6, pady=6)
        self.txtB = ctk.CTkTextbox(body, width=500, height=400); self.txtB.pack(side="left", fill="both", expand=True, padx=6, pady=6)

        btns = ctk.CTkFrame(page, fg_color="transparent"); btns.pack(fill="x", pady=4)
        ctk.CTkButton(btns, text="Carregar Arquivo A", command=lambda:self._load_file(self.txtA)).pack(side="left", padx=4)
        ctk.CTkButton(btns, text="Carregar Arquivo B", command=lambda:self._load_file(self.txtB)).pack(side="left", padx=4)
        ctk.CTkButton(btns, text="Comparar", command=self._compare_texts).pack(side="left", padx=4)
        ctk.CTkButton(btns, text="Exportar DOCX", command=self._export_diff_report).pack(side="left", padx=4)

        self.preview = ctk.CTkTextbox(page, height=120); self.preview.pack(fill="x", padx=12, pady=6)
        navbtns = ctk.CTkFrame(page, fg_color="transparent"); navbtns.pack(fill="x", pady=(0,8))
        ctk.CTkButton(navbtns, text="⬆ Anterior", command=self._go_to_prev_diff).pack(side="left", padx=4)
        ctk.CTkButton(navbtns, text="⬇ Próxima", command=self._go_to_next_diff).pack(side="left", padx=4)

    def _load_file(self, textbox):
        path = filedialog.askopenfilename(filetypes=[("Text/Word", "*.txt *.docx")])
        if not path: return
        text = ""
        if path.endswith(".docx"):
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
        else:
            text = Path(path).read_text(encoding="utf-8")
        textbox.delete("1.0","end")
        textbox.insert("1.0", text)

    
    
    def _compare_texts(self):
        a = self.txtA.get("1.0","end-1c").splitlines()
        b = self.txtB.get("1.0","end-1c").splitlines()
        sm = difflib.SequenceMatcher(None, a, b)

        # Limpar tags anteriores
        for tag in ["replace","delete","insert"]:
            self.txtA.tag_remove(tag,"1.0","end")
            self.txtB.tag_remove(tag,"1.0","end")

        # Definir estilos de destaque (cores suaves tipo WinMerge)
        self.txtA.tag_config("replace", background="#FFF2CC", foreground="black")
        self.txtB.tag_config("replace", background="#F0DDA8", foreground="black")
        self.txtA.tag_config("delete", background="#F4CCCC", foreground="black")
        self.txtB.tag_config("insert", background="#D9EAD3", foreground="black")

        # Guardar blocos de diferenças para navegação
        self.diffs = []
        self.preview.delete("1.0","end")
        for tag,i1,i2,j1,j2 in sm.get_opcodes():
            if tag=="equal":
                continue
            self.diffs.append((tag,i1,i2,j1,j2))
            self.preview.insert("end", f"{tag.upper()} A[{i1+1}:{i2}] ↔ B[{j1+1}:{j2}]")
            if tag=="replace":
                self.txtA.tag_add("replace", f"{i1+1}.0", f"{i2}.end")
                self.txtB.tag_add("replace", f"{j1+1}.0", f"{j2}.end")
            elif tag=="delete":
                self.txtA.tag_add("delete", f"{i1+1}.0", f"{i2}.end")
            elif tag=="insert":
                self.txtB.tag_add("insert", f"{j1+1}.0", f"{j2}.end")
        self.current_diff = -1

    def _go_to_next_diff(self):
        if not hasattr(self,"diffs") or not self.diffs: return
        self.current_diff = (self.current_diff+1) % len(self.diffs)
        self._scroll_to_diff(self.current_diff)

    def _go_to_prev_diff(self):
        if not hasattr(self,"diffs") or not self.diffs: return
        self.current_diff = (self.current_diff-1) % len(self.diffs)
        self._scroll_to_diff(self.current_diff)

    def _scroll_to_diff(self, idx):
        tag,i1,i2,j1,j2 = self.diffs[idx]
        if tag in ("replace","delete"):
            self.txtA.see(f"{i1+1}.0")
        if tag in ("replace","insert"):
            self.txtB.see(f"{j1+1}.0")

    
    def _export_diff_report(self):
        from docx import Document as DocxDoc
        from docx.shared import RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        doc = DocxDoc()
        a = self.txtA.get("1.0","end-1c").splitlines()
        b = self.txtB.get("1.0","end-1c").splitlines()
        sm = difflib.SequenceMatcher(None,a,b)
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text, table.rows[0].cells[1].text = "Texto A","Texto B"

        for tag,i1,i2,j1,j2 in sm.get_opcodes():
            maxlen = max(i2-i1,j2-j1)
            for k in range(maxlen):
                ra = a[i1+k] if i1+k<i2 else ""
                rb = b[j1+k] if j1+k<j2 else ""
                row = table.add_row().cells
                row[0].text, row[1].text = ra, rb
                # aplicar cor de fundo suave no DOCX
                if tag=="replace":
                    for c in row:
                        c._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FFF2CC"/>'.format(nsdecls('w'))))
                elif tag=="delete":
                    row[0]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="F4CCCC"/>'.format(nsdecls('w'))))
                elif tag=="insert":
                    row[1]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w'))))

        save = filedialog.asksaveasfilename(defaultextension=".docx")
        if save:
            doc.save(save)
            messagebox.showinfo("Exportar","Relatório exportado!")

    def _page_pdf2word(self):
        page = ctk.CTkFrame(self.content, fg_color="transparent")
        self.pages["pdf2word"] = page
        header = ctk.CTkFrame(page, fg_color=self.cores["bg_card"], height=62)
        header.pack(fill="x", padx=16, pady=(16,0)); header.pack_propagate(False)
        ctk.CTkLabel(header, text="📄 Conversor PDF → Word", font=ctk.CTkFont(size=20, weight="bold")).pack(side="left", padx=18, pady=16)

        body = ctk.CTkFrame(page, fg_color=self.cores["bg_card"]); body.pack(fill="both", expand=True, padx=16, pady=12)
        self.pdf_in = tk.StringVar(); self.pdf_out = tk.StringVar()
        row1 = ctk.CTkFrame(body, fg_color="transparent"); row1.pack(fill="x", pady=4)
        ctk.CTkEntry(row1, textvariable=self.pdf_in).pack(side="left", fill="x", expand=True, padx=4)
        ctk.CTkButton(row1, text="Selecionar PDF", command=self._pick_pdf).pack(side="left", padx=4)

        row2 = ctk.CTkFrame(body, fg_color="transparent"); row2.pack(fill="x", pady=4)
        ctk.CTkEntry(row2, textvariable=self.pdf_out).pack(side="left", fill="x", expand=True, padx=4)
        ctk.CTkButton(row2, text="Salvar como Word", command=self._save_pdf2word).pack(side="left", padx=4)

        ctk.CTkButton(body, text="Converter", command=self._convert_pdf2word).pack(pady=10)

    def _pick_pdf(self):
        p = filedialog.askopenfilename(filetypes=[("PDF","*.pdf")])
        if p: self.pdf_in.set(p)

    def _save_pdf2word(self):
        p = filedialog.asksaveasfilename(defaultextension=".docx")
        if p: self.pdf_out.set(p)

    def _convert_pdf2word(self):
        if not self.pdf_in.get() or not self.pdf_out.get():
            messagebox.showwarning("Aviso","Selecione entrada e saída"); return
        if not PDFPLUMBER_OK:
            messagebox.showerror("Erro","pdfplumber não instalado"); return
        doc = Document()
        import pdfplumber
        with pdfplumber.open(self.pdf_in.get()) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                doc.add_paragraph(text)
        doc.save(self.pdf_out.get())
        messagebox.showinfo("Sucesso","PDF convertido para Word!")


    # ---------- Página: WINMERGE (Comparação estilo WinMerge) ----------
    def _page_winmerge(self):
        page = ctk.CTkFrame(self.content, fg_color="transparent")
        self.pages["winmerge"] = page
        header = ctk.CTkFrame(page, fg_color=self.cores["bg_card"], height=62)
        header.pack(fill="x", padx=16, pady=(16,0)); header.pack_propagate(False)
        ctk.CTkLabel(header, text="🔀 Comparação estilo WinMerge", font=ctk.CTkFont(size=20, weight="bold")).pack(side="left", padx=18, pady=16)

        body = ctk.CTkFrame(page, fg_color=self.cores["bg_card"]); body.pack(fill="both", expand=True, padx=16, pady=12)
        self.txtA = ctk.CTkTextbox(body, width=500, height=400); self.txtA.pack(side="left", fill="both", expand=True, padx=6, pady=6)
        self.txtB = ctk.CTkTextbox(body, width=500, height=400); self.txtB.pack(side="left", fill="both", expand=True, padx=6, pady=6)

        btns = ctk.CTkFrame(page, fg_color="transparent"); btns.pack(fill="x", pady=4)
        ctk.CTkButton(btns, text="Carregar Arquivo A", command=lambda:self._load_file(self.txtA)).pack(side="left", padx=4)
        ctk.CTkButton(btns, text="Carregar Arquivo B", command=lambda:self._load_file(self.txtB)).pack(side="left", padx=4)
        ctk.CTkButton(btns, text="Comparar", command=self._compare_texts).pack(side="left", padx=4)
        ctk.CTkButton(btns, text="Exportar DOCX", command=self._export_diff_report).pack(side="left", padx=4)

        self.preview = ctk.CTkTextbox(page, height=120); self.preview.pack(fill="x", padx=12, pady=6)
        navbtns = ctk.CTkFrame(page, fg_color="transparent"); navbtns.pack(fill="x", pady=(0,8))
        ctk.CTkButton(navbtns, text="⬆ Anterior", command=self._go_to_prev_diff).pack(side="left", padx=4)
        ctk.CTkButton(navbtns, text="⬇ Próxima", command=self._go_to_next_diff).pack(side="left", padx=4)

    def _load_file(self, textbox):
        path = filedialog.askopenfilename(filetypes=[("Text/Word", "*.txt *.docx")])
        if not path: return
        if path.endswith(".docx"):
            from docx import Document
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
        else:
            text = Path(path).read_text(encoding="utf-8")
        textbox.delete("1.0","end")
        textbox.insert("1.0", text)

    def _compare_texts(self):
        from difflib import Differ
        a_lines = self.txtA.get("1.0","end-1c").splitlines()
        b_lines = self.txtB.get("1.0","end-1c").splitlines()
        differ = Differ()
        diff = list(differ.compare(a_lines, b_lines))

        # limpar tags
        for tag in ["replace","delete","insert"]: 
            self.txtA.tag_remove(tag,"1.0","end")
            self.txtB.tag_remove(tag,"1.0","end")

        # configurar cores suaves estilo WinMerge
        self.txtA.tag_config("replace", background="#FFF2CC", foreground="black")
        self.txtB.tag_config("replace", background="#FFF2CC", foreground="black")
        self.txtA.tag_config("delete", background="#F4CCCC", foreground="black")
        self.txtB.tag_config("insert", background="#D9EAD3", foreground="black")

        self.preview.delete("1.0","end")
        self.diffs = []
        lineA = 1
        lineB = 1
        for d in diff:
            code = d[:2]; text = d[2:]
            if code == "  ":
                lineA += 1; lineB += 1
            elif code == "- ":
                self.txtA.tag_add("delete", f"{lineA}.0", f"{lineA}.end")
                self.preview.insert("end", f"- {text}\n")
                self.diffs.append(("delete",lineA,None))
                lineA += 1
            elif code == "+ ":
                self.txtB.tag_add("insert", f"{lineB}.0", f"{lineB}.end")
                self.preview.insert("end", f"+ {text}\n")
                self.diffs.append(("insert",None,lineB))
                lineB += 1
            elif code == "? ":
                self.preview.insert("end", f"~ {text}\n")

        self.current_diff = -1

    def _go_to_next_diff(self):
        if not hasattr(self,"diffs") or not self.diffs: return
        self.current_diff = (self.current_diff+1) % len(self.diffs)
        self._scroll_to_diff(self.current_diff)

    def _go_to_prev_diff(self):
        if not hasattr(self,"diffs") or not self.diffs: return
        self.current_diff = (self.current_diff-1) % len(self.diffs)
        self._scroll_to_diff(self.current_diff)

    def _scroll_to_diff(self, idx):
        tag,i,j = self.diffs[idx]
        if tag=="delete" and i is not None:
            self.txtA.see(f"{i}.0")
        elif tag=="insert" and j is not None:
            self.txtB.see(f"{j}.0")

    def _export_diff_report(self):
        from docx import Document as DocxDoc
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from difflib import Differ

        a_lines = self.txtA.get("1.0","end-1c").splitlines()
        b_lines = self.txtB.get("1.0","end-1c").splitlines()
        differ = Differ()
        diff = list(differ.compare(a_lines, b_lines))

        doc = DocxDoc()
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text, table.rows[0].cells[1].text = "Texto A","Texto B"

        for d in diff:
            code = d[:2]; text = d[2:]
            row = table.add_row().cells
            if code=="  ":
                row[0].text, row[1].text = text, text
            elif code=="- ":
                row[0].text = text
                row[0]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="F4CCCC"/>'.format(nsdecls('w'))))
            elif code=="+ ":
                row[1].text = text
                row[1]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w'))))
            elif code=="? ":
                row[0].text = text
                row[1].text = text
                for c in row:
                    c._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FFF2CC"/>'.format(nsdecls('w'))))

        save = filedialog.asksaveasfilename(defaultextension=".docx")
        if save:
            doc.save(save)
            messagebox.showinfo("Exportar","Relatório exportado!")

    # ---------- Página: PDF2WORD ----------
    def _page_pdf2word(self):
        page = ctk.CTkFrame(self.content, fg_color="transparent")
        self.pages["pdf2word"] = page
        header = ctk.CTkFrame(page, fg_color=self.cores["bg_card"], height=62)
        header.pack(fill="x", padx=16, pady=(16,0)); header.pack_propagate(False)
        ctk.CTkLabel(header, text="📄 Conversor PDF → Word", font=ctk.CTkFont(size=20, weight="bold")).pack(side="left", padx=18, pady=16)

        body = ctk.CTkFrame(page, fg_color=self.cores["bg_card"]); body.pack(fill="both", expand=True, padx=16, pady=12)
        self.pdf_in = tk.StringVar(); self.pdf_out = tk.StringVar()
        row1 = ctk.CTkFrame(body, fg_color="transparent"); row1.pack(fill="x", pady=4)
        ctk.CTkEntry(row1, textvariable=self.pdf_in).pack(side="left", fill="x", expand=True, padx=4)
        ctk.CTkButton(row1, text="Selecionar PDF", command=self._pick_pdf).pack(side="left", padx=4)

        row2 = ctk.CTkFrame(body, fg_color="transparent"); row2.pack(fill="x", pady=4)
        ctk.CTkEntry(row2, textvariable=self.pdf_out).pack(side="left", fill="x", expand=True, padx=4)
        ctk.CTkButton(row2, text="Salvar como Word", command=self._save_pdf2word).pack(side="left", padx=4)

        ctk.CTkButton(body, text="Converter", command=self._convert_pdf2word).pack(pady=10)

    def _pick_pdf(self):
        p = filedialog.askopenfilename(filetypes=[("PDF","*.pdf")])
        if p: self.pdf_in.set(p)

    def _save_pdf2word(self):
        p = filedialog.asksaveasfilename(defaultextension=".docx")
        if p: self.pdf_out.set(p)

    def _convert_pdf2word(self):
        if not self.pdf_in.get() or not self.pdf_out.get():
            messagebox.showwarning("Aviso","Selecione entrada e saída"); return
        try:
            import pdfplumber
        except:
            messagebox.showerror("Erro","pdfplumber não instalado"); return
        doc = Document()
        with pdfplumber.open(self.pdf_in.get()) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                doc.add_paragraph(text)
        doc.save(self.pdf_out.get())
        messagebox.showinfo("Sucesso","PDF convertido para Word!")

    def _page_sobre(self):
        page = ctk.CTkFrame(self.content, fg_color="transparent")
        self.pages["sobre"] = page
        header = ctk.CTkFrame(page, fg_color=self.cores["bg_card"], height=62); header.pack(fill="x", padx=16, pady=(16, 0)); header.pack_propagate(False)
        ctk.CTkLabel(header, text="ℹ️ Sobre", font=ctk.CTkFont(size=20, weight="bold")).pack(side="left", padx=18, pady=16)
        body = ctk.CTkFrame(page, fg_color=self.cores["bg_card"]); body.pack(fill="both", expand=True, padx=16, pady=12)
        ctk.CTkLabel(body, text=f"{APP_NAME} v{APP_VERSION}\nBuild single-file com melhorias não-invasivas.\n© 2025", text_color=self.cores["text_secondary"]).pack(pady=16)

    # ---------- Helpers Preview/Lote ----------
    def _page_kpi_placeholder(self): pass

# -------- CLI/Headless --------
def run_headless(word_path: str, out_path: str, modelo: str=None, apply_norm=False, gloss_rules:str=""):
    if not os.path.exists(word_path) or not word_path.lower().endswith(".docx"):
        raise FileNotFoundError("Word .docx inválido.")
    doc = Document(word_path)
    texto = "\n".join(p.text for p in doc.paragraphs)
    padrao = r"(Template\s*Name:.*?(?=Template\s*Name:|$))"
    blocos = re.findall(padrao, texto, re.DOTALL | re.IGNORECASE)

    if not modelo:
        nome = "Modelo de templates v1.1.xlsx"
        home = Path.home()
        candidatos = [
            Path.cwd() / nome,
            Path.cwd().parent / nome,
            home / nome,
            home / "Desktop" / nome,
            home / "Documents" / nome,
            home / "Downloads" / nome,
        ]
        for c in candidatos:
            if c.is_file():
                modelo = str(c); break
    if not modelo or not os.path.exists(modelo):
        raise FileNotFoundError("Modelo de templates não encontrado.")

    wb = openpyxl.load_workbook(modelo)
    ws = wb.active
    headers = [c.value for c in ws[1] if c.value]
    if ws.max_row > 1: ws.delete_rows(2, ws.max_row - 1)

    rules = {}
    for line in (gloss_rules or "").splitlines():
        if "=" in line:
            k,v = line.split("=",1)
            rules[k.strip()] = v.strip()

    def formatar_prioridade(valor):
        m = re.match(r"(\d)", valor or "")
        mapa = {"1": "1 - Critical", "2": "2 - High", "3": "3 - Medium", "4": "4 - Low"}
        return mapa.get(m.group(1) if m else "", "3 - Medium")

    def extrair_campos(bloco: str, ci_global: str):
        campos = {
            "Name": "",
            "Short description": "",
            "Short Description INC": "",
            "Configuration item": "",
            "Category": "",
            "Subcategory": "",
            "Impact": "",
            "Urgency": "",
            "Assignment group": "",
            "Description": "",
            "State": "Open",
            "Fill assigned to with the current user?": "",
            "Active": "TRUE",
            "Global": "FALSE",
            "Groups": "ITIL Processes Knowledge",
            "Article": "",
        }
        for linha in (bloco or "").splitlines():
            if not linha.strip(): continue
            low = linha.lower()
            if low.startswith("template name:"):
                name = linha.split(":", 1)[1].strip()
                campos["Name"] = name
                up = name.upper()
                if up.startswith("L1"):
                    campos["State"] = "Resolved"
                    campos["Fill assigned to with the current user?"] = "TRUE"
                elif up.startswith("L2"):
                    campos["State"] = "Open"
                    campos["Fill assigned to with the current user?"] = "FALSE"
            elif low.startswith("short description:"):
                sd = linha.split(":", 1)[1].strip()
                campos["Short description"] = sd
                campos["Short Description INC"] = sd
            elif re.match(r"^ci\s*:", low):
                ci = linha.split(":", 1)[1].strip()
                campos["Configuration item"] = ci
            elif low.startswith("category:"):
                campos["Category"] = linha.split(":", 1)[1].strip()
            elif low.startswith("subcategory:"):
                campos["Subcategory"] = linha.split(":", 1)[1].strip()
            elif low.startswith("impact:"):
                campos["Impact"] = formatar_prioridade(linha.split(":", 1)[1].strip())
            elif low.startswith("urgency:"):
                campos["Urgency"] = formatar_prioridade(linha.split(":", 1)[1].strip())
            elif low.startswith("assignment group:"):
                campos["Assignment group"] = linha.split(":", 1)[1].strip()
            elif low.startswith("description:"):
                campos["Description"] = linha.split(":", 1)[1].strip()
            elif "kb00" in low:
                campos["Article"] = linha.strip()
        if apply_norm and rules:
            campos = apply_normalizers(campos, True, rules)
        return campos

    # CI global (opcional)
    texto_all = "\n".join(p.text for p in doc.paragraphs)
    mt_ci = re.findall(r"\bCI\s*:\s*(.+)", texto_all, re.IGNORECASE)
    ci_global = ", ".join(v.strip() for v in mt_ci if v.strip())

    for i, bloco in enumerate(blocos, start=1):
        campos = extrair_campos(bloco, ci_global)
        row = i + 1
        for col_idx, header in enumerate(headers, start=1):
            ws.cell(row=row, column=col_idx, value=campos.get(header, ""))

    out_dir = Path(out_path).parent
    out_dir.mkdir(parents=True, exist_ok=True)
    wb.save(out_path)
    write_manifest(inputs=[word_path, modelo], outputs=[out_path], schema_version="2.0.0", app_version=APP_VERSION, notes={"mode":"cli"})
    return out_path

# -------- Main --------
def main():
    ensure_admin_exists()
    # CLI
    parser = argparse.ArgumentParser(add_help=False)
    parser.add_argument("--headless", action="store_true")
    parser.add_argument("--word", type=str, help=".docx de entrada")
    parser.add_argument("--out", type=str, help=".xlsx de saída")
    parser.add_argument("--modelo", type=str, help="modelo .xlsx")
    parser.add_argument("--apply-norm", action="store_true")
    parser.add_argument("--gloss", type=str, default="", help="regras termo=valor, separadas por \\n")
    parser.add_argument("--help", action="store_true")
    args, _ = parser.parse_known_args()

    if args.help:
        print("AGC CLI:\n --headless --word IN.docx --out OUT.xlsx [--modelo MODEL.xlsx] [--apply-norm] [--gloss 'kg=kg\\nlitro=L']")
        return

    if args.headless:
        if not args.word or not args.out:
            print("Parâmetros obrigatórios: --word e --out")
            sys.exit(2)
        try:
            out = run_headless(args.word, args.out, args.modelo, args.apply_norm, args.gloss)
            print(f"OK: {out}")
            return
        except Exception as e:
            print(f"ERRO: {e}")
            sys.exit(1)

    # GUI
    login = LoginWindow(); login.mainloop()
    user = login.current_user or {"username":"guest","name":"Convidado","role":"user"}
    AGCApp(user)

if __name__ == "__main__":
    main()



